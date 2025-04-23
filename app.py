import cv2
import numpy as np
import pytesseract
import streamlit as st
from streamlit_drawable_canvas import st_canvas
from docx import Document
import os
from PIL import Image

# Set a modern, full-screen layout and the title
st.set_page_config(layout="wide")
st.title("🚀 Smart Document Scanner and OCR")

# Custom CSS for a more modern look
st.markdown("""
    <style>
        .stApp {
            background: linear-gradient(135deg, #ff9e2a, #ff4e00);
            color: white;
        }
        .stButton>button {
            background-color: #ff4e00;
            color: white;
            font-weight: bold;
        }
        .stTextArea>textarea {
            background-color: #f7f7f7;
            color: #333;
            border: 2px solid #ff4e00;
        }
        .stHeader {
            font-size: 2.5rem;
            font-weight: 800;
        }
    </style>
""", unsafe_allow_html=True)

# Define helper functions
def plot_image(img, title="Image"):
    st.image(cv2.cvtColor(img, cv2.COLOR_BGR2RGB), caption=title, use_column_width=True)

def order_points(pts):
    rect = np.zeros((4, 2), dtype="float32")
    s = pts.sum(axis=1)
    diff = np.diff(pts, axis=1)
    rect[0] = pts[np.argmin(s)]  # Top-left
    rect[2] = pts[np.argmax(s)]  # Bottom-right
    rect[1] = pts[np.argmin(diff)]  # Top-right
    rect[3] = pts[np.argmax(diff)]  # Bottom-left
    return rect

def four_point_transform(image, pts):
    rect = order_points(pts)
    (tl, tr, br, bl) = rect
    widthA = np.linalg.norm(br - bl)
    widthB = np.linalg.norm(tr - tl)
    maxWidth = int(max(widthA, widthB))
    heightA = np.linalg.norm(tr - br)
    heightB = np.linalg.norm(tl - bl)
    maxHeight = int(max(heightA, heightB))
    dst = np.array([
        [0, 0],
        [maxWidth-1, 0],
        [maxWidth-1, maxHeight-1],
        [0, maxHeight-1]
    ], dtype="float32")
    M = cv2.getPerspectiveTransform(rect, dst)
    warped = cv2.warpPerspective(image, M, (maxWidth, maxHeight))
    return warped

def extract_text(image):
    return pytesseract.image_to_string(image)

# File uploader
uploaded_file = st.file_uploader("📤 Upload your image here!", type=["png", "jpg", "jpeg", "bmp", "tiff"])

if uploaded_file is not None:
    file_bytes = np.asarray(bytearray(uploaded_file.read()), dtype=np.uint8)
    image = cv2.imdecode(file_bytes, cv2.IMREAD_COLOR)

    orig_h, orig_w = image.shape[:2]
    display_w = 800
    scale_factor = display_w / orig_w
    display_h = int(orig_h * scale_factor)
    resized_image = cv2.resize(image, (display_w, display_h))

    st.subheader("🎨 Draw a rectangle to select the portion of the document you want to scan!")
    canvas_result = st_canvas(
        fill_color="rgba(255, 255, 255, 0.3)",
        stroke_width=3,
        background_image=Image.fromarray(cv2.cvtColor(resized_image, cv2.COLOR_BGR2RGB)),
        update_streamlit=True,
        height=display_h,
        width=display_w,
        drawing_mode="rect",
        key="canvas"
    )

    if canvas_result.json_data and "objects" in canvas_result.json_data:
        rect_data = canvas_result.json_data["objects"]
        if len(rect_data) > 0:
            rect = rect_data[0]  # Get the first drawn rectangle (in case multiple rectangles are drawn)
            if "left" in rect and "top" in rect and "width" in rect and "height" in rect:
                # Capture the rectangle's coordinates
                left = rect["left"]
                top = rect["top"]
                width = rect["width"]
                height = rect["height"]
                
                # Convert the rectangle coordinates back to the original image size
                left = int(left / scale_factor)
                top = int(top / scale_factor)
                width = int(width / scale_factor)
                height = int(height / scale_factor)

                # Crop the image based on the rectangle
                cropped_image = image[top:top+height, left:left+width]

                st.subheader("📄 Cropped Image")
                plot_image(cropped_image)

                # Perform perspective transformation on the cropped image
                scanned = four_point_transform(resized_image, np.array([
                    [left, top], 
                    [left + width, top], 
                    [left + width, top + height], 
                    [left, top + height]
                ], dtype="float32"))

                st.subheader("🖼️ Scanned Image")
                plot_image(scanned)

                # Extract text from the scanned image
                st.subheader("🔍 Extracted Text")
                text = extract_text(scanned)
                st.text_area("Extracted Text", text, height=200)

                if st.button("💾 Save as Word Document"):
                    image_name = os.path.splitext(os.path.basename(uploaded_file.name))[0]
                    doc = Document()
                    doc.add_heading(f'Scanned Document: {image_name}', 0)
                    doc.add_paragraph(text)
                    output_path = f"{image_name}_scanned.docx"
                    doc.save(output_path)
                    st.success(f"✅ Document saved as: {output_path}")
            else:
                st.warning("⚠️ Please draw a rectangle to select the region.")
