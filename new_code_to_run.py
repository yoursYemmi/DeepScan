import cv2
import numpy as np
import pytesseract
import matplotlib.pyplot as plt
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import os
import tempfile

# Setup Tesseract path if needed
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def plot_image(img, title="Image"):
    plt.imshow(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
    plt.title(title)
    plt.xticks([]); plt.yticks([])
    plt.show()

def get_four_points_from_user(image):
    points = []

    def click_event(event, x, y, flags, param):
        if event == cv2.EVENT_LBUTTONDOWN and len(points) < 4:
            points.append((x, y))
            cv2.circle(image, (x, y), 5, (0, 255, 0), -1)
            cv2.imshow("Select 4 corners (TL, TR, BR, BL)", image)

    cv2.imshow("Select 4 corners (TL, TR, BR, BL)", image)
    cv2.setMouseCallback("Select 4 corners (TL, TR, BR, BL)", click_event)

    while True:
        key = cv2.waitKey(1) & 0xFF
        if len(points) == 4 or key == 27:
            break

    cv2.destroyAllWindows()
    return np.array(points, dtype="float32")

def order_points(pts):
    rect = np.zeros((4, 2), dtype="float32")
    s = pts.sum(axis=1)
    diff = np.diff(pts, axis=1)
    rect[0] = pts[np.argmin(s)]  # Top-left
    rect[2] = pts[np.argmax(s)]  # Bottom-right
    rect[1] = pts[np.argmin(diff)]  # Top-right
    rect[3] = pts[np.argmax(diff)]  # Bottom-left
    return rect

def four_point_transform(image, pts):
    rect = order_points(pts)
    (tl, tr, br, bl) = rect
    widthA = np.linalg.norm(br - bl)
    widthB = np.linalg.norm(tr - tl)
    maxWidth = int(max(widthA, widthB))

    heightA = np.linalg.norm(tr - br)
    heightB = np.linalg.norm(tl - bl)
    maxHeight = int(max(heightA, heightB))

    dst = np.array([
        [0, 0],
        [maxWidth-1, 0],
        [maxWidth-1, maxHeight-1],
        [0, maxHeight-1]
    ], dtype="float32")

    M = cv2.getPerspectiveTransform(rect, dst)
    warped = cv2.warpPerspective(image, M, (maxWidth, maxHeight))
    return warped

def preprocess_image(img):
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    resized = cv2.resize(gray, (2667, 1500), interpolation=cv2.INTER_LINEAR)
    return resized

def extract_text(image):
    return pytesseract.image_to_string(image)

# --- GUI File Selection ---
root = tk.Tk()
root.withdraw()
file_path = filedialog.askopenfilename(title="Select an image", filetypes=[("Images", "*.png *.jpg *.jpeg *.bmp *.tiff")])
if not file_path:
    print("No file selected.")
    exit()

# --- Handle format conversion ---
temp_png_path = None
img = cv2.imread(file_path)
if not file_path.lower().endswith('.png'):
    temp_png_path = os.path.join(tempfile.gettempdir(), "temp_converted.png")
    cv2.imwrite(temp_png_path, img)
    file_path = temp_png_path
    img = cv2.imread(file_path)

# --- Show original image ---
plot_image(img, "Original Image")

# --- Resize for OCR and User Selection ---
resized_img = cv2.resize(img, (2667, 1500), interpolation=cv2.INTER_LINEAR)

# --- Let user select contour points ---
points = get_four_points_from_user(resized_img.copy())
if len(points) != 4:
    print("Not enough points selected. Exiting.")
    exit()

# --- Transform and show scanned image ---
scanned = four_point_transform(resized_img, points)
plot_image(scanned, "Scanned Image")

# --- Preprocess scanned image ---
processed = preprocess_image(scanned)

# --- OCR ---
text = extract_text(processed)
print("Extracted Text:\n", text)



# --- Save Word Document (if approved) ---
should_save = messagebox.askyesno("Save Document", "Do you want to save the scanned content as a Word document?")
if should_save:
    image_name = os.path.splitext(os.path.basename(file_path))[0]
    doc = Document()
    doc.add_heading(f'Scanned Document: {image_name}', 0)
    doc.add_paragraph(text)
    output_path = f"{image_name}_scanned.docx"  # <-- Updated line
    doc.save(output_path)
    print(f"Document saved as: {output_path}")
else:
    print("Document not saved.")

